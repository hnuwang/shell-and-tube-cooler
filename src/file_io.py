from __future__ import annotations

import json
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook

from config import DesignConfig

PARAMETER_LABELS: dict[str, str] = {
    "hot_mass_flow_kg_s": "煤油质量流量 (kg/s)",
    "hot_inlet_temp_c": "煤油入口温度 (℃)",
    "hot_outlet_temp_c": "煤油出口温度 (℃)",
    "cold_inlet_temp_c": "冷却水入口温度 (℃)",
    "cold_outlet_temp_c": "冷却水出口温度 (℃)",
    "hot_pressure_mpa": "煤油压力 (MPa)",
    "cold_pressure_mpa": "冷却水压力 (MPa)",
    "shell_passes": "壳程数",
    "tube_passes": "管程数",
    "tube_outer_diameter_m": "管外径 (m)",
    "tube_inner_diameter_m": "管内径 (m)",
    "tube_length_candidates_m": "管长候选 (m)",
    "pitch_ratio": "管间距比",
    "layout_angle_deg": "布管角 (deg)",
    "shell_clearance_m": "壳体单边间隙 (m)",
    "baffle_spacing_ratio": "挡板间距比",
    "baffle_spacing_min_m": "挡板间距最小值 (m)",
    "baffle_spacing_max_m": "挡板间距最大值 (m)",
    "baffle_cut_ratio": "挡板切口率",
    "tube_wall_thermal_conductivity_w_m_k": "管壁导热系数 (W/m/K)",
    "fouling_resistance_tube_m2_k_w": "管侧污垢热阻 (m2·K/W)",
    "fouling_resistance_shell_m2_k_w": "壳侧污垢热阻 (m2·K/W)",
    "initial_overall_u_w_m2_k": "初选总传热系数 U (W/m2/K)",
    "use_u_recalculation": "启用 U 校核复算",
    "tube_velocity_min_m_s": "管程最小流速 (m/s)",
    "tube_velocity_max_m_s": "管程最大流速 (m/s)",
    "shell_velocity_min_m_s": "壳程最小流速 (m/s)",
    "shell_velocity_max_m_s": "壳程最大流速 (m/s)",
    "allowable_tube_pressure_drop_pa": "许用管程压降 (Pa)",
    "allowable_shell_pressure_drop_pa": "许用壳程压降 (Pa)",
    "area_margin_min": "面积裕量最小值",
    "area_margin_max": "面积裕量最大值",
    "max_iterations": "最大迭代次数",
    "strict_property_range": "物性越界严格报错",
    "print_intermediate": "命令行打印中间量",
    "export_markdown_tables": "启用 Markdown 表输出",
    "tube_inlet_loss_coefficient": "入口局部阻力系数",
    "tube_return_loss_coefficient_per_pass": "回弯局部阻力系数",
    "tube_outlet_loss_coefficient": "出口局部阻力系数",
    "shell_friction_factor_constant": "壳程摩擦系数常数",
}

LABEL_TO_KEY = {label: key for key, label in PARAMETER_LABELS.items()}


def load_config_from_json(path: str | Path) -> DesignConfig:
    """Load a configuration mapping from JSON."""
    data = json.loads(Path(path).read_text(encoding="utf-8"))
    return DesignConfig.from_mapping(data)


def load_config_from_excel(path: str | Path) -> DesignConfig:
    """Load configuration from an Excel sheet with key/value rows."""
    workbook = load_workbook(Path(path), data_only=True)
    worksheet = workbook[workbook.sheetnames[0]]
    values: dict[str, Any] = {}
    for row in worksheet.iter_rows(min_row=1, max_col=3, values_only=True):
        raw_key = row[0]
        raw_value = row[1]
        if raw_key is None or raw_value is None:
            continue
        key_text = str(raw_key).strip()
        if key_text.lower() in {"parameter", "参数", "key", "field"}:
            continue
        key = LABEL_TO_KEY.get(key_text, key_text)
        values[key] = raw_value
    if not values:
        raise ValueError("Excel 参数表未读取到有效键值，请检查首列为参数名、次列为参数值。")
    return DesignConfig.from_mapping(values)


def export_word_report(path: str | Path, design) -> None:
    """Export a concise Word report with summary tables."""
    document = Document()
    title = document.add_heading("固定管板式管壳式煤油冷却器计算结果", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = document.add_paragraph()
    summary.add_run("综合结论：").bold = True
    summary.add_run(
        f"热负荷 {design.thermal_result.heat_duty_w:.2f} W，冷却水流量 {design.thermal_result.cold_mass_flow_kg_s:.3f} kg/s，"
        f"选定管长 {design.mechanical_result.tube_geometry.tube_length_m:.2f} m，总管数 {design.mechanical_result.tube_geometry.tube_count} 根，"
        f"管程压降 {design.hydraulic_result.tube_pressure_drop_pa:.1f} Pa，壳程压降 {design.hydraulic_result.shell_pressure_drop_pa:.1f} Pa。"
    )

    for title_text, key in (
        ("输入参数表", "input"),
        ("物性参数表", "properties"),
        ("传热计算结果表", "thermal"),
        ("结构计算结果表", "mechanical"),
        ("阻力计算结果表", "hydraulic"),
    ):
        document.add_heading(title_text, level=1)
        rows = design.result_tables[key]
        headers = list(rows[0].keys())
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = str(header)
        for row in rows:
            cells = table.add_row().cells
            for index, header in enumerate(headers):
                value = row[header]
                if isinstance(value, float):
                    cells[index].text = f"{value:.6g}"
                else:
                    cells[index].text = str(value)

    document.save(str(path))


def export_excel_report(path: str | Path, design, workspace_dir: str | Path) -> None:
    """Export the design result tables to Excel via the local workbook builder."""
    workspace_path = Path(workspace_dir)
    scripts_dir = workspace_path / "scripts"
    scripts_dir.mkdir(exist_ok=True)
    temp_json = scripts_dir / "_design_export_payload.json"
    payload = {
        "result_tables": design.result_tables,
        "summary": {
            "heat_duty_w": design.thermal_result.heat_duty_w,
            "cold_mass_flow_kg_s": design.thermal_result.cold_mass_flow_kg_s,
            "required_area_m2": design.thermal_result.required_area_m2,
            "actual_area_m2": design.mechanical_result.actual_area_m2,
            "tube_pressure_drop_pa": design.hydraulic_result.tube_pressure_drop_pa,
            "shell_pressure_drop_pa": design.hydraulic_result.shell_pressure_drop_pa,
            "design_state": "可行",
        },
    }
    temp_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    node_executable = Path.home() / ".cache" / "codex-runtimes" / "codex-primary-runtime" / "dependencies" / "node" / "bin" / "node.exe"
    script_path = workspace_path / "scripts" / "export_results_workbook.mjs"
    command = [str(node_executable), str(script_path), str(temp_json), str(Path(path))]
    subprocess.run(command, check=True, cwd=str(workspace_path))


def build_parameter_rows(config: DesignConfig) -> list[tuple[str, Any]]:
    """Build ordered rows for parameter preview/export."""
    rows: list[tuple[str, Any]] = []
    config_dict = config.to_dict()
    config_dict.pop("base_dir", None)
    for key, label in PARAMETER_LABELS.items():
        value = config_dict[key]
        if isinstance(value, list):
            value = ", ".join(str(item) for item in value)
        rows.append((label, value))
    return rows


def write_parameter_template_excel(path: str | Path, config: DesignConfig) -> None:
    """Write a simple Excel parameter template using openpyxl."""
    from openpyxl import Workbook

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "参数表"
    worksheet["A1"] = "参数"
    worksheet["B1"] = "参数值"
    worksheet["C1"] = "字段名"
    for index, (label, value) in enumerate(build_parameter_rows(config), start=2):
        key = list(PARAMETER_LABELS.keys())[index - 2]
        worksheet.cell(row=index, column=1, value=label)
        worksheet.cell(row=index, column=2, value=value)
        worksheet.cell(row=index, column=3, value=key)
    workbook.save(path)
